import os
import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

CONFIG_FILE = "config.json"

def main():
    with open(CONFIG_FILE) as f:
        config = json.load(f)

    ROOT_DIR = config["root_dir"]
    OUTPUT_DIR = config["output_dir"]
    OUTPUT_FILE_NAME = config["output_file_name"]
    OUTPUT_FILE = f"{OUTPUT_DIR}/{OUTPUT_FILE_NAME}.docx"
    CODE_FONT = config["code_font"]
    HEADING = config["heading"]
    EXTENSIONS = config.get("include_extensions", [])

    doc = Document()

    os.makedirs(ROOT_DIR, exist_ok=True)
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    
    title = doc.add_heading(HEADING)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Traversing the file structure
    for folder, _, files in os.walk(ROOT_DIR): # (folder_path, subdirectories, filenames)
        for filename in files:
            if any(filename.endswith(ext) for ext in EXTENSIONS):
                file_path = os.path.join(folder, filename)
                rel_path = os.path.relpath(file_path, ROOT_DIR)

                # Add Section Header
                doc.add_heading(f"File: {rel_path}", level=2)

                # Adding the code into the output file
                with open(file_path, "r") as f:
                    content = f.read()

                para = doc.add_paragraph(content)
                para.style.font.name = CODE_FONT
                run = para.runs
                for r in run:
                    r.font.size = Pt(10)
                
    doc.add_section()

    # Add a paragraph at the end
    CREDITS_FILE = ".credits.json"
    with open(CREDITS_FILE) as f:
        credits = json.load(f)
    
    PROJECT_NAME = credits["project_name"]
    GITHUB_LINK = credits["github_link"]

    end_footer = doc.sections[-1].footer
    end_footer_para = end_footer.paragraphs[0]
    end_footer_para.text = f"© 2025 Anshul Badhani – {PROJECT_NAME} – {GITHUB_LINK}"
    end_footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    end_footer_para.runs[0].font.size = Pt(10)
    # Save to DOCX
    doc.save(OUTPUT_FILE)
    print(f"✅ Assignment saved to {OUTPUT_FILE}")

if __name__ == "__main__":
    main()
